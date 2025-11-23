#!/usr/bin/env python3
"""
Extract structure from Word document templates.
"""

from docx import Document
import sys

def extract_doc_structure(file_path):
    """Extract text and structure from a Word document."""
    try:
        doc = Document(file_path)
        print(f"\n{'='*80}")
        print(f"Document: {file_path}")
        print(f"{'='*80}\n")

        print(f"Total Paragraphs: {len(doc.paragraphs)}")
        print(f"Total Tables: {len(doc.tables)}")
        print(f"\n{'='*80}")
        print("DOCUMENT CONTENT:")
        print(f"{'='*80}\n")

        for i, para in enumerate(doc.paragraphs):
            style = para.style.name
            text = para.text.strip()
            if text:
                print(f"[{i:03d}] [{style}] {text}")

        if doc.tables:
            print(f"\n{'='*80}")
            print("TABLES:")
            print(f"{'='*80}\n")
            for i, table in enumerate(doc.tables):
                print(f"\nTable {i+1}:")
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    print(f"  {' | '.join(row_text)}")

        return True
    except Exception as e:
        print(f"Error reading {file_path}: {e}")
        return False

if __name__ == "__main__":
    # Try to read all template documents
    files = [
        r"College\67F_B Tech Semester VII Project Report (1) (1).doc",
        r"College\Activity 2_AirGroove.docx"
    ]

    for file in files:
        extract_doc_structure(file)
        print("\n" + "="*80 + "\n")
